"""Generate the MLOps Final Report as a .docx file (~10 pages)."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ---- Styles ----
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0, 51, 102)

# ---- Title Page ----
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('MLOps Assignment 1\n')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('End-to-End ML Model Development,\nCI/CD, and Production Deployment')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

course = doc.add_paragraph()
course.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = course.add_run('Course: MLOps (S2-25_AMLCSZG523)')
run.font.size = Pt(12)

doc.add_paragraph()

# Team table
team_para = doc.add_paragraph()
team_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = team_para.add_run('Team: Group 5')
run.font.size = Pt(14)
run.bold = True

table = doc.add_table(rows=5, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Light Shading Accent 1'
hdr = table.rows[0]
hdr.cells[0].text = 'ID'
hdr.cells[1].text = 'Name'
members = [
    ('2025cs05043', 'Tushar Chouhan'),
    ('2025cs05010', 'Kelli L Prasanna Kumar'),
    ('2025cs05045', 'Jyoti Chugh'),
    ('2025cs05047', 'Madhan M'),
]
for i, (sid, name) in enumerate(members):
    table.rows[i + 1].cells[0].text = sid
    table.rows[i + 1].cells[1].text = name

doc.add_page_break()

# ---- Table of Contents ----
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Introduction',
    '2. Data Acquisition & Exploratory Data Analysis',
    '3. Feature Engineering & Model Development',
    '4. Experiment Tracking with MLflow',
    '5. Model Packaging & Reproducibility',
    '6. CI/CD Pipeline & Automated Testing',
    '7. Model Containerization',
    '8. Production Deployment (Minikube + Podman)',
    '9. Monitoring & Logging (Prometheus + Grafana)',
    '10. Architecture Overview',
    '11. Conclusion',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ---- 1. Introduction ----
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    'This report documents the end-to-end development and deployment of a machine learning '
    'solution for predicting heart disease risk using the UCI Heart Disease dataset. The project '
    'follows modern MLOps best practices, encompassing data acquisition, exploratory data analysis, '
    'feature engineering, model development and evaluation, experiment tracking, model packaging, '
    'automated testing with CI/CD pipelines, containerization, production deployment on Kubernetes, '
    'and real-time monitoring.'
)
doc.add_paragraph(
    'The primary objective is to build a binary classifier that predicts the presence or absence '
    'of heart disease based on 13 patient health features (age, sex, blood pressure, cholesterol, '
    'etc.). The trained model is served via a FastAPI REST API, containerized with Docker, deployed '
    'to a local Kubernetes cluster (Minikube with Podman as the container runtime), and monitored '
    'using Prometheus and Grafana.'
)
doc.add_heading('1.1 Dataset Overview', level=2)
doc.add_paragraph(
    'The Heart Disease dataset from the UCI Machine Learning Repository (ID: 45) contains 303 '
    'patient records with 13 clinical features and a target variable. The original target has '
    'values 0-4, which we convert to binary (0 = no disease, 1 = disease present). After cleaning, '
    'the dataset has 297 usable samples with no missing values.'
)
doc.add_heading('1.2 Technology Stack', level=2)
tech_items = [
    'Python 3.11, scikit-learn, pandas, numpy, matplotlib, seaborn',
    'MLflow for experiment tracking',
    'FastAPI + Uvicorn for model serving',
    'Docker / Podman for containerization',
    'Minikube + Podman driver for Kubernetes deployment',
    'Prometheus + Grafana for monitoring',
    'GitHub Actions for CI/CD',
    'pytest for unit testing',
]
for t in tech_items:
    doc.add_paragraph(t, style='List Bullet')

doc.add_page_break()

# ---- 2. Data Acquisition & EDA ----
doc.add_heading('2. Data Acquisition & Exploratory Data Analysis', level=1)
doc.add_paragraph(
    'The dataset is fetched programmatically using the ucimlrepo Python library, which provides '
    'direct access to UCI datasets. This ensures reproducibility — the same data can be re-fetched '
    'at any time without manual download steps.'
)
doc.add_heading('2.1 Data Cleaning', level=2)
doc.add_paragraph(
    'Data preprocessing steps include: (1) converting all columns to numeric types using '
    'pd.to_numeric with coerce for non-parseable values, (2) filling missing values with column '
    'medians for numeric features and mode for categorical features, (3) dropping any remaining '
    'null rows after coercion, and (4) converting the multi-class target to binary (presence vs. '
    'absence of heart disease). The cleaned dataset is saved as heart_disease_cleaned.csv for '
    'downstream reproducibility.'
)
doc.add_heading('2.2 Exploratory Data Analysis', level=2)
doc.add_paragraph(
    'We produced four categories of professional visualizations:'
)
eda_items = [
    'Class Distribution: Bar chart and pie chart showing the target balance (~46% disease, '
    '~54% no disease), confirming the dataset is reasonably balanced and does not require '
    'aggressive resampling techniques.',
    'Feature Distributions: Histograms for all 13 features, colored by target class, revealing '
    'which features show visible separation between disease and no-disease groups. Notable '
    'separations were observed in thal, cp, ca, and oldpeak.',
    'Correlation Heatmap: A triangular heatmap of Pearson correlations, highlighting that '
    'thal (+0.53), ca (+0.47), cp (+0.43), and exang (+0.44) have the strongest associations '
    'with the target variable.',
    'Box Plots: Side-by-side box plots for key continuous features (age, trestbps, chol, '
    'thalach, oldpeak), showing differences in distributions between disease classes and '
    'identifying potential outliers in cholesterol and blood pressure.',
]
for item in eda_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    'All plots are saved to the screenshots/ folder for inclusion in this report and the '
    'GitHub repository.'
)

# Insert screenshots if available
screenshot_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'screenshots')
for img_name, caption in [
    ('class_distribution.png', 'Figure 1: Class Distribution'),
    ('correlation_heatmap.png', 'Figure 2: Correlation Heatmap'),
    ('cv_comparison.png', 'Figure 3: Cross-Validation Model Comparison'),
    ('roc_curves.png', 'Figure 4: ROC Curves'),
    ('confusion_matrices.png', 'Figure 5: Confusion Matrices'),
]:
    img_path = os.path.join(screenshot_dir, img_name)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)

doc.add_page_break()

# ---- 3. Feature Engineering & Model Development ----
doc.add_heading('3. Feature Engineering & Model Development', level=1)
doc.add_heading('3.1 Feature Preparation', level=2)
doc.add_paragraph(
    'All 13 features are retained without dimensionality reduction, as the dataset is small and '
    'all features carry clinical relevance. Feature scaling is performed using StandardScaler, '
    'which transforms each feature to have zero mean and unit variance. This is critical for '
    'Logistic Regression (which is sensitive to feature scales) and generally beneficial for '
    'gradient-based algorithms.'
)
doc.add_paragraph(
    'The dataset is split into 80% training and 20% test sets using stratified sampling '
    '(test_size=0.2, stratify=y) to preserve class proportions in both subsets.'
)

doc.add_heading('3.2 Model Training', level=2)
doc.add_paragraph('Three classification models were trained and compared:')
models_desc = [
    ('Logistic Regression', 'A linear model with L2 regularization (max_iter=1000). '
     'Serves as a strong interpretable baseline.'),
    ('Random Forest', 'An ensemble of 100 decision trees with bootstrap aggregation. '
     'Handles non-linear relationships and provides feature importance rankings.'),
    ('Gradient Boosting', 'A sequential ensemble of 100 weak learners that minimizes '
     'a differentiable loss function. Typically achieves the highest accuracy on structured data.'),
]
for name, desc in models_desc:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('3.3 Evaluation', level=2)
doc.add_paragraph(
    'Models are evaluated using 5-fold stratified cross-validation on the training set, and '
    'final performance is measured on the held-out test set. Metrics include accuracy, precision, '
    'recall, F1-score, and ROC-AUC. The best model is selected based on the highest test-set '
    'ROC-AUC score, balancing overall discrimination ability across all classification thresholds.'
)

doc.add_page_break()

# ---- 4. Experiment Tracking ----
doc.add_heading('4. Experiment Tracking with MLflow', level=1)
doc.add_paragraph(
    'MLflow is integrated to track all experiment runs. The tracking URI is set to a local '
    'file store (file:./mlruns) under a single experiment named "Heart_Disease_Classification". '
    'For each model run, we log:'
)
mlflow_items = [
    'All hyperparameters (retrieved via model.get_params())',
    'Five test-set metrics: accuracy, precision, recall, F1-score, ROC-AUC',
    'Five cross-validation metrics with mean and standard deviation',
    'The trained scikit-learn model as an MLflow artifact',
    'Visualization plots (ROC curves, confusion matrices, CV comparison)',
]
for item in mlflow_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    'A dedicated "Best Model" run logs the final scikit-learn Pipeline (StandardScaler + '
    'best classifier), the pickle file, metadata JSON, and the cleaned dataset for full '
    'reproducibility. The MLflow dashboard can be launched with: mlflow ui --backend-store-uri file:./mlruns'
)

# ---- 5. Model Packaging ----
doc.add_heading('5. Model Packaging & Reproducibility', level=1)
doc.add_paragraph(
    'The final model is packaged as a scikit-learn Pipeline object that encapsulates both the '
    'StandardScaler preprocessing and the classifier in a single object. This means any new input '
    'data only needs to be passed through pipeline.predict() — no separate scaling step is required.'
)
doc.add_paragraph('Three artifacts ensure full reproducibility:')
packaging_items = [
    'model_pipeline.pkl: The complete Pipeline saved via pickle. Can be loaded and used for '
    'inference with zero additional setup.',
    'model_metadata.json: Contains feature names, model type, training sample count, and '
    'test-set accuracy/ROC-AUC. Used by the API to validate input features at runtime.',
    'requirements.txt: Pinned versions of all Python dependencies (fastapi==0.115.0, '
    'scikit-learn==1.5.0, mlflow==2.16.0, etc.) ensuring environment reproducibility.',
]
for item in packaging_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ---- 6. CI/CD ----
doc.add_heading('6. CI/CD Pipeline & Automated Testing', level=1)
doc.add_heading('6.1 Unit Tests', level=2)
doc.add_paragraph(
    'The test suite (test/test_model.py) contains three test classes with a total of 10+ test cases:'
)
test_items = [
    'TestDataProcessing: Tests dataset download, shape validation, binary target conversion, '
    'missing value handling, and feature scaling correctness.',
    'TestModel: Tests model training, prediction format, probability output validity, pipeline '
    'consistency (deterministic predictions), and model serialization (save/load round-trip).',
    'TestAPI: Tests input data validation and prediction response format.',
]
for item in test_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.2 GitHub Actions Pipeline', level=2)
doc.add_paragraph(
    'The CI/CD pipeline (.github/workflows/ci-cd.yml) is triggered on every push and pull request '
    'to the main branch. It consists of three sequential jobs:'
)
cicd_items = [
    'lint-and-test: Installs dependencies, runs flake8 linting on app.py and test/, '
    'then runs pytest with JUnit XML output. Test results are uploaded as workflow artifacts.',
    'train-model: After tests pass, the model is retrained from scratch using the UCI dataset. '
    'The trained model_pipeline.pkl and model_metadata.json are uploaded as artifacts.',
    'build-docker: Downloads model artifacts, builds the Docker image, runs the container, '
    'waits 10 seconds for startup, then smoke-tests the /health and /predict endpoints.',
]
for item in cicd_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    'The pipeline enforces that code linting, all unit tests, model training, Docker build, '
    'and endpoint smoke tests must all pass before a merge is allowed. Any failure produces '
    'clear error logs visible in the GitHub Actions UI.'
)

# ---- 7. Containerization ----
doc.add_heading('7. Model Containerization', level=1)
doc.add_paragraph(
    'The model is served via a FastAPI application (app.py) that exposes the following endpoints:'
)
endpoint_table = doc.add_table(rows=6, cols=3)
endpoint_table.style = 'Light Shading Accent 1'
headers = ['Endpoint', 'Method', 'Description']
for i, h in enumerate(headers):
    endpoint_table.rows[0].cells[i].text = h
endpoints = [
    ('/', 'GET', 'Root welcome message'),
    ('/health', 'GET', 'Health check (used by k8s probes)'),
    ('/predict', 'POST', 'Make prediction (JSON input, returns label + confidence)'),
    ('/metrics', 'GET', 'Prometheus metrics endpoint'),
    ('/model-info', 'GET', 'Model metadata (type, accuracy, features)'),
]
for i, (ep, method, desc) in enumerate(endpoints):
    endpoint_table.rows[i + 1].cells[0].text = ep
    endpoint_table.rows[i + 1].cells[1].text = method
    endpoint_table.rows[i + 1].cells[2].text = desc

doc.add_paragraph()
doc.add_paragraph(
    'The Dockerfile uses python:3.11-slim as the base image, installs dependencies from '
    'requirements.txt, copies the application code and model artifacts, exposes port 8000, '
    'configures a health check, and runs uvicorn. The image can be built and tested with either '
    'Docker or Podman.'
)

doc.add_page_break()

# ---- 8. Production Deployment ----
doc.add_heading('8. Production Deployment (Minikube + Podman)', level=1)
doc.add_paragraph(
    'For production deployment, we use Minikube with Podman as the container runtime driver. '
    'Podman is a daemonless, rootless container engine that provides a Docker-compatible CLI, '
    'making it suitable for development and CI environments where running a Docker daemon is '
    'not desirable.'
)
doc.add_heading('8.1 Kubernetes Resources', level=2)
doc.add_paragraph('The deployment manifest (deployment/k8s-deployment.yaml) defines:')
k8s_items = [
    'Deployment: 2 replicas of the heart-disease-api container, with resource requests '
    '(256Mi memory, 250m CPU) and limits (512Mi memory, 500m CPU).',
    'Liveness Probe: HTTP GET to /health every 30 seconds with a 15-second initial delay. '
    'Kubernetes will restart the pod if the health check fails.',
    'Readiness Probe: HTTP GET to /health every 10 seconds with a 10-second initial delay. '
    'The pod will not receive traffic until it passes the readiness check.',
    'Service: A LoadBalancer service that routes external traffic on port 80 to the container '
    'port 8000.',
]
for item in k8s_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('8.2 Deployment Process', level=2)
doc.add_paragraph(
    'A deployment script (deploy_minikube.sh) automates the entire process: starting minikube '
    'with the podman driver, building the Docker image inside minikube\'s container runtime, '
    'applying the Kubernetes manifests, and waiting for all pods to reach Ready status. '
    'The API can then be accessed via minikube service or kubectl port-forward.'
)

# ---- 9. Monitoring ----
doc.add_heading('9. Monitoring & Logging (Prometheus + Grafana)', level=1)
doc.add_heading('9.1 Application Metrics', level=2)
doc.add_paragraph(
    'The FastAPI application exports Prometheus metrics via the /metrics endpoint using the '
    'prometheus_client library. Three custom metrics are defined:'
)
metrics_items = [
    'prediction_requests_total (Counter): Total prediction requests, labeled by status '
    '(success/error). Tracks API reliability.',
    'prediction_latency_seconds (Histogram): Prediction latency distribution. '
    'Enables percentile analysis (p50, p90, p99).',
    'prediction_results (Counter): Prediction outcome counts, labeled by result '
    '(Disease/No Disease). Monitors model behavior drift.',
]
for item in metrics_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('9.2 Application Logging', level=2)
doc.add_paragraph(
    'All API requests are logged using Python\'s logging module. Logs are written to both '
    'stdout (for container log aggregation) and a file (api.log). Each prediction log entry '
    'includes the prediction result, confidence score, and latency. Errors are logged at '
    'ERROR level with full exception details.'
)

doc.add_heading('9.3 Monitoring Stack', level=2)
doc.add_paragraph(
    'The docker-compose.yml defines a three-service monitoring stack:'
)
stack_items = [
    'API (port 8000): The FastAPI application with Prometheus metrics.',
    'Prometheus (port 9090): Scrapes the /metrics endpoint every 15 seconds. '
    'Configuration is provided via deployment/prometheus.yml.',
    'Grafana (port 3000): Pre-provisioned with a Prometheus datasource and a custom dashboard. '
    'The dashboard includes panels for request counts, latency percentiles, request rate, '
    'prediction outcome trends, and a prediction distribution pie chart.',
]
for item in stack_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    'Grafana dashboard provisioning is fully automated — on startup, Grafana loads the datasource '
    'and dashboard JSON from mounted volumes, requiring no manual configuration.'
)

doc.add_page_break()

# ---- 10. Architecture ----
doc.add_heading('10. Architecture Overview', level=1)
doc.add_paragraph(
    'The following diagram illustrates the end-to-end architecture of the MLOps pipeline:'
)
arch_text = """
┌─────────────────┐     ┌──────────────────┐     ┌──────────────┐
│  UCI Repository  │────>│  Jupyter Notebook │────>│    MLflow     │
│  (Data Source)   │     │  (EDA + Training) │     │   Tracking    │
└─────────────────┘     └────────┬─────────┘     └──────────────┘
                                 │
                          ┌──────▼───────┐
                          │ model_pipeline│
                          │   .pkl        │
                          └──────┬───────┘
                                 │
                    ┌────────────▼────────────┐
                    │  FastAPI App (app.py)   │
                    │  /predict  /health      │
                    │  /metrics  /model-info  │
                    └────────────┬────────────┘
                                 │
                    ┌────────────▼────────────┐
                    │   Docker / Podman       │
                    │   Container Image       │
                    └────────────┬────────────┘
                                 │
              ┌──────────────────▼──────────────────┐
              │   Minikube (Podman Driver)           │
              │   Deployment: 2 replicas             │
              │   Service: LoadBalancer (port 80)    │
              └──────────────────┬──────────────────┘
                                 │
              ┌──────────────────▼──────────────────┐
              │   Prometheus + Grafana               │
              │   (Metrics & Dashboard)              │
              └─────────────────────────────────────┘

    CI/CD: GitHub Actions
    ┌────────────────────────────────────────┐
    │  Lint → Test → Train → Build → Deploy  │
    └────────────────────────────────────────┘
"""
p = doc.add_paragraph()
run = p.add_run(arch_text)
run.font.name = 'Courier New'
run.font.size = Pt(8)

doc.add_page_break()

# ---- 11. Conclusion ----
doc.add_heading('11. Conclusion', level=1)
doc.add_paragraph(
    'This project demonstrates a complete, production-ready MLOps pipeline for heart disease '
    'prediction. Starting from raw data acquisition through the UCI repository, we performed '
    'thorough exploratory data analysis, trained and evaluated three classification models, '
    'and tracked all experiments with MLflow for full auditability.'
)
doc.add_paragraph(
    'The best-performing model was packaged as a scikit-learn Pipeline and served via a FastAPI '
    'REST API with Swagger documentation. The application was containerized using Docker (compatible '
    'with Podman), with a multi-stage CI/CD pipeline in GitHub Actions that enforces code quality '
    'through linting, comprehensive unit tests, model retraining, and container smoke tests.'
)
doc.add_paragraph(
    'For production deployment, we utilized Minikube with Podman as the container runtime, '
    'deploying the API as a Kubernetes Deployment with 2 replicas, health probes, resource limits, '
    'and a LoadBalancer service. Real-time monitoring was implemented using Prometheus for metrics '
    'collection and Grafana for visualization, with auto-provisioned dashboards tracking prediction '
    'request counts, latency percentiles, error rates, and prediction outcome distributions.'
)
doc.add_paragraph(
    'Key achievements of this project include:'
)
conclusions = [
    'End-to-end automation: From data ingestion to model serving, every step is scripted and reproducible.',
    'Experiment tracking: All model runs are logged in MLflow with parameters, metrics, and artifacts.',
    'Automated quality gates: CI/CD pipeline enforces linting, testing, and smoke testing before deployment.',
    'Container-native deployment: Docker/Podman containerization with Kubernetes orchestration.',
    'Observable system: Prometheus metrics and Grafana dashboards provide real-time visibility into API health.',
    'Production best practices: Health probes, resource limits, graceful degradation, and structured logging.',
]
for item in conclusions:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph(
    'The project successfully mirrors real-world production MLOps workflows and demonstrates '
    'proficiency in the full ML lifecycle — from experimentation to production deployment and monitoring.'
)

# ---- Save ----
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'MLOps_Final_Report.docx')
doc.save(output_path)
print(f'Report saved to: {output_path}')
